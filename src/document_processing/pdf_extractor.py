import os
import re
from typing import Any, Dict, List
from config.logger import logger
from src.core.exceptions import DocumentProcessingException

try:
    import fitz  # PyMuPDF
    _FITZ_AVAILABLE = True
except Exception as _e_fitz:
    logger.warning(f"PyMuPDF fitz import unavailable: {_e_fitz}")
    _FITZ_AVAILABLE = False

try:
    import pdfplumber
    _PDFPLUMBER_AVAILABLE = True
except Exception as _e_plumber:
    logger.warning(f"pdfplumber import unavailable: {_e_plumber}")
    _PDFPLUMBER_AVAILABLE = False

try:
    from pypdf import PdfReader
    _PYPDF_AVAILABLE = True
except Exception as _e_pypdf:
    logger.warning(f"pypdf import unavailable: {_e_pypdf}")
    _PYPDF_AVAILABLE = False

try:
    import docx
    _DOCX_AVAILABLE = True
except Exception:
    _DOCX_AVAILABLE = False


class ExtractedPage:
    def __init__(self, page_number: int, text: str, word_count: int, font_names: List[str] = None):
        self.page_number = page_number
        self.text = text
        self.word_count = word_count
        self.font_names = font_names or []


class ExtractedDocument:
    def __init__(self, title: str, pages: List[ExtractedPage], metadata: Dict[str, Any], total_pages: int, total_words: int):
        self.title = title
        self.pages = pages
        self.metadata = metadata
        self.total_pages = total_pages
        self.total_words = total_words


class UniversalDocumentExtractor:
    """Universal Multi-Format Document Extractor for PDFs, Word, Code, Data, Text & Any File Type."""

    def extract(self, file_path: str, filename: str) -> ExtractedDocument:
        if not os.path.exists(file_path):
            raise DocumentProcessingException(f"File not found: {file_path}")

        ext = os.path.splitext(filename)[1].lower()

        # PDF Documents
        if ext == ".pdf":
            return self._extract_pdf(file_path, filename)

        # Word Documents
        elif ext in [".docx", ".doc"]:
            return self._extract_docx(file_path, filename)

        # Plain Text, Markdown, Code & Structured Data Files
        else:
            return self._extract_text_or_generic(file_path, filename)

    def _extract_pdf(self, file_path: str, filename: str) -> ExtractedDocument:
        if _FITZ_AVAILABLE:
            try:
                return self._extract_pymupdf(file_path, filename)
            except Exception as e1:
                logger.warning(f"PyMuPDF extraction failed for {filename}: {e1}. Trying pdfplumber fallback.")

        if _PDFPLUMBER_AVAILABLE:
            try:
                return self._extract_pdfplumber(file_path, filename)
            except Exception as e2:
                logger.warning(f"pdfplumber extraction failed: {e2}. Trying pypdf fallback.")

        if _PYPDF_AVAILABLE:
            try:
                return self._extract_pypdf(file_path, filename)
            except Exception as e3:
                logger.error(f"pypdf extraction failed: {e3}")

        # Final Fallback: generic string extraction from PDF binary stream
        return self._extract_text_or_generic(file_path, filename)

    def _extract_pymupdf(self, file_path: str, filename: str) -> ExtractedDocument:
        doc = fitz.open(file_path)
        pages: List[ExtractedPage] = []
        total_words = 0
        raw_metadata = doc.metadata or {}

        metadata = {
            "author": raw_metadata.get("author", ""),
            "title": raw_metadata.get("title", "") or filename,
            "subject": raw_metadata.get("subject", ""),
            "creator": raw_metadata.get("creator", ""),
            "producer": raw_metadata.get("producer", ""),
            "format": doc.format,
            "is_encrypted": doc.is_encrypted
        }

        for i in range(len(doc)):
            page = doc[i]
            text = page.get_text("text") or ""
            cleaned_text = " ".join(text.split())
            words = len(cleaned_text.split())
            total_words += words

            pages.append(ExtractedPage(
                page_number=i + 1,
                text=cleaned_text,
                word_count=words
            ))

        doc.close()
        return ExtractedDocument(
            title=metadata["title"],
            pages=pages,
            metadata=metadata,
            total_pages=len(pages),
            total_words=total_words
        )

    def _extract_pdfplumber(self, file_path: str, filename: str) -> ExtractedDocument:
        pages: List[ExtractedPage] = []
        total_words = 0

        with pdfplumber.open(file_path) as pdf:
            metadata = {
                "author": pdf.metadata.get("Author", "") if pdf.metadata else "",
                "title": pdf.metadata.get("Title", "") or filename if pdf.metadata else filename,
                "producer": pdf.metadata.get("Producer", "") if pdf.metadata else ""
            }

            for idx, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                cleaned_text = " ".join(text.split())
                words = len(cleaned_text.split())
                total_words += words

                pages.append(ExtractedPage(
                    page_number=idx + 1,
                    text=cleaned_text,
                    word_count=words
                ))

        return ExtractedDocument(
            title=metadata["title"],
            pages=pages,
            metadata=metadata,
            total_pages=len(pages),
            total_words=total_words
        )

    def _extract_pypdf(self, file_path: str, filename: str) -> ExtractedDocument:
        reader = PdfReader(file_path)
        pages: List[ExtractedPage] = []
        total_words = 0

        meta = reader.metadata or {}
        metadata = {
            "author": str(meta.get("/Author", "")),
            "title": str(meta.get("/Title", "")) or filename,
        }

        for idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            cleaned_text = " ".join(text.split())
            words = len(cleaned_text.split())
            total_words += words

            pages.append(ExtractedPage(
                page_number=idx + 1,
                text=cleaned_text,
                word_count=words
            ))

        return ExtractedDocument(
            title=metadata["title"],
            pages=pages,
            metadata=metadata,
            total_pages=len(pages),
            total_words=total_words
        )

    def _extract_docx(self, file_path: str, filename: str) -> ExtractedDocument:
        if _DOCX_AVAILABLE:
            try:
                doc = docx.Document(file_path)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                cleaned_text = " ".join(full_text.split())
                words = len(cleaned_text.split())

                pages = [ExtractedPage(page_number=1, text=cleaned_text, word_count=words)]
                return ExtractedDocument(
                    title=os.path.splitext(filename)[0].replace("_", " ").title(),
                    pages=pages,
                    metadata={"format": "Word Document (.docx)"},
                    total_pages=1,
                    total_words=words
                )
            except Exception as e:
                logger.warning(f"python-docx failed for {filename}: {e}. Falling back to text reader.")

        return self._extract_text_or_generic(file_path, filename)

    def _extract_text_or_generic(self, file_path: str, filename: str) -> ExtractedDocument:
        raw_text = ""
        # Try UTF-8 -> Latin-1 -> UTF-16
        for enc in ["utf-8", "latin-1", "utf-16", "cp1252"]:
            try:
                with open(file_path, "r", encoding=enc, errors="ignore") as f:
                    raw_text = f.read()
                if raw_text:
                    break
            except Exception:
                continue

        # If binary data, extract printable strings
        if not raw_text or len(raw_text.strip()) == 0:
            with open(file_path, "rb") as f:
                content_bytes = f.read()
                # Find ASCII printable strings
                printable_strings = re.findall(b"[a-zA-Z0-9\\s\\.,;:\\-_/\\\\(\\)\\[\\]\\{\\}]{4,}", content_bytes)
                raw_text = " ".join([s.decode("ascii", errors="ignore") for s in printable_strings])

        cleaned_text = " ".join(raw_text.split())
        words = len(cleaned_text.split())

        # Page pagination (approx 400 words per logical page)
        chunk_size = 2000
        text_chunks = [cleaned_text[i:i + chunk_size] for i in range(0, max(1, len(cleaned_text)), chunk_size)]

        pages: List[ExtractedPage] = []
        for idx, chunk in enumerate(text_chunks, start=1):
            w_count = len(chunk.split())
            pages.append(ExtractedPage(
                page_number=idx,
                text=chunk,
                word_count=w_count
            ))

        return ExtractedDocument(
            title=os.path.splitext(filename)[0].replace("_", " ").title(),
            pages=pages,
            metadata={"format": f"Universal File ({os.path.splitext(filename)[1]})"},
            total_pages=len(pages),
            total_words=words
        )


pdf_extractor = UniversalDocumentExtractor()
